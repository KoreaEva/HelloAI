from docx import Document
from openpyxl import load_workbook
from docx2pdf import convert
import os
import win32com.client

wdFormatPDF = 17
# for i, paragraph in enumerate(doc.paragraphs):
#     print(str(i) + ": " + paragraph.text)

def UpdateDocx(doument_number, name, birthday, duration, course, year, month, day):
    doc = Document('수료증템플릿.docx')
    doc.paragraphs[0].text = doc.paragraphs[0].text.replace('<?document_number?>',document_number)
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace('<?name?>',name)
    doc.paragraphs[4].text = doc.paragraphs[4].text.replace('<?birthday?>',str(birthday.date()))
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace('<?duration?>',duration)
    doc.paragraphs[8].text = doc.paragraphs[8].text.replace('<?course?>',course)
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace('<?year?>',str(year))
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace('<?month?>',str(month))
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace('<?day?>',str(day))

    print('- {}님의 수료증을 생성하고 있습니다.'.format(name))

    filename = '수료증_{0}.'.format(name)
    doc.save(filename + 'docx')

    # pdf = aw.Document(filename + 'docx')
    # pdf.save(filename + '.pdf')
    print('- {}님의 PDF를 생성하고 있습니다.'.format(name))

    inputFile = os.path.abspath(filename + 'docx')
    outputFile = os.path.abspath(filename + 'pdf')
    word = win32com.client.Dispatch("Word.Application")
    docx = word.Documents.Open(inputFile)
    docx.SaveAs(outputFile, FileFormat=wdFormatPDF)
    word.Quit()

wb = load_workbook('수강생 명단.xlsx')
ws = wb.active

for i in range(ws.max_row):

    document_number = ws.cell(row=i+1, column=1).value
    name = ws.cell(row=i+1, column=2).value
    birthday = ws.cell(row=i+1, column=3).value
    duration = ws.cell(row=i+1, column=4).value
    course = ws.cell(row=i+1, column=5).value
    year = ws.cell(row=i+1, column=6).value
    month = ws.cell(row=i+1, column=7).value
    day = ws.cell(row=i+1, column=8).value

    if i != 0:
        print('[{}/{}]: {}님의 데이터를 처리하고 있습니다.'.format(ws.max_row-1, i, name ))

        UpdateDocx(doument_number=document_number,
                name=name,
                birthday=birthday,
                duration=duration,
                course=course,
                year=year,
                month=month,
                day=day)